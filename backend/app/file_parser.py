"""
File parser for resume extraction from PDF, DOCX, and plain text files.
"""

import io
import logging
from typing import Union

logger = logging.getLogger(__name__)


def parse_resume_file(file_content: bytes, file_extension: str) -> str:
    """
    Parse resume file and extract text.
    
    Supports: PDF, DOCX, TXT
    """
    
    file_extension = file_extension.lower()
    
    if file_extension == 'pdf':
        return parse_pdf(file_content)
    elif file_extension in ['docx', 'doc']:
        return parse_docx(file_content)
    elif file_extension == 'txt':
        return parse_text(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def parse_pdf(file_content: bytes) -> str:
    """Parse PDF file and extract text."""
    try:
        import pdfplumber
        
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
        
        return text.strip()
    
    except ImportError:
        logger.error("pdfplumber not installed")
        raise ValueError("PDF parsing requires pdfplumber. Install with: pip install pdfplumber")
    except Exception as e:
        logger.error(f"PDF parsing error: {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_content: bytes) -> str:
    """Parse DOCX file and extract text."""
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except ImportError:
        logger.error("python-docx not installed")
        raise ValueError("DOCX parsing requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX parsing error: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_text(file_content: Union[bytes, str]) -> str:
    """Parse plain text file."""
    try:
        if isinstance(file_content, bytes):
            return file_content.decode('utf-8', errors='ignore').strip()
        else:
            return file_content.strip()
    except Exception as e:
        logger.error(f"Text parsing error: {str(e)}")
        raise ValueError(f"Failed to parse text: {str(e)}")
